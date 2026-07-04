from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/SOP_真实机器人AI控制系统.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, text, bold=True)
        set_cell_shading(cell, "F2F4F7")
        cell.width = widths[idx]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text)
            cells[idx].width = widths[idx]
    doc.add_paragraph()
    return table


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SO-ARM101 苹果采摘真实机器人 AI 控制系统 SOP")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("policy_runtime_service + robot_client 闭环执行架构").italic = True

    doc.add_heading("1. 标准声明", level=1)
    doc.add_paragraph(
        "本系统为真实机器人AI控制系统：控制链采用 policy_runtime_service + robot_client 的闭环执行架构。"
        "mock_task 仅用于UI兜底显示，不参与任何机器人控制。"
    )

    doc.add_heading("2. 实时控制链路", level=1)
    add_code_block(
        doc,
        [
            "frontend",
            "  -> FastAPI backend",
            "  -> task_control",
            "  -> policy_runtime_service",
            "  -> LeRobot ACT policy inference",
            "  -> robot_client",
            "  -> SO-ARM101 robot execution",
        ],
    )
    add_bullets(
        doc,
        [
            "policy_runtime_service 是唯一 AI 推理入口，负责加载并运行 ACT policy。",
            "robot_client 是唯一机械臂硬件执行入口，负责读取观测并发送动作。",
            "frontend 只发送任务意图和模式选择，不直接生成或发送机器人动作。",
            "task_control 负责任务生命周期编排，不能绕过 policy_runtime_service。",
        ],
    )

    doc.add_heading("3. 苹果成熟度分级控制", level=1)
    doc.add_paragraph(
        "系统使用统一字段 target_maturity 控制目标苹果成熟度。yellow apple 表示半成熟果控制策略。"
    )
    add_table(
        doc,
        ["UI 选项", "后端字段", "Policy input"],
        [
            ["成熟果（红苹果）", 'target_maturity: "red"', 'target_maturity="red"'],
            ["半成熟果（黄苹果）", 'target_maturity: "yellow"', 'target_maturity="yellow"'],
        ],
        [Inches(1.8), Inches(2.2), Inches(2.2)],
    )
    add_bullets(
        doc,
        [
            "task_control 不再使用旧的“只红/多摘/半摘”逻辑。",
            "local 和 remote 模式均将同一个 target_maturity 传入 policy_runtime_service。",
            "policy_runtime_service 在推理循环中把 target_maturity 注入 policy input。",
            "robot_client 接口保持不变。",
        ],
    )

    doc.add_heading("4. UI 展示链路", level=1)
    add_code_block(
        doc,
        [
            "robot_status + backend_state + optional mock fallback",
            "  -> status_fusion",
            "  -> websocket / HTTP",
            "  -> frontend UI",
        ],
    )
    doc.add_paragraph(
        "展示链只影响页面可读性，不改变控制逻辑。mock_task 只允许为 UI fallback 提供展示数据。"
    )

    doc.add_heading("5. local / remote 部署定义", level=1)
    add_table(
        doc,
        ["模式", "部署链路", "说明"],
        [
            [
                "local",
                "backend -> policy_runtime_service（本机进程） -> robot_client -> SO-ARM101",
                "本机部署真实 AI 控制闭环。",
            ],
            [
                "remote",
                "backend -> HTTP -> Robot PC -> policy_runtime_service（远程进程） -> robot_client -> SO-ARM101",
                "Robot PC 承载远程 policy runtime 和 robot client。",
            ],
        ],
        [Inches(1.0), Inches(3.7), Inches(1.8)],
    )

    doc.add_heading("6. 操作流程", level=1)
    add_bullets(
        doc,
        [
            "启动 FastAPI backend 和前端 UI。",
            "确认 LEROBOT_POLICY_MODEL_PATH、POLICY_INFERENCE_HZ 和机器人连接配置。",
            "前端选择 target_maturity（red 或 yellow），并通过 /set_target_apple 写入后端状态。",
            "前端调用 /start_task，task_control 根据 mode 调用本机或远程 policy_runtime_service。",
            "policy_runtime_service 读取 robot_client observation，执行 ACT policy 推理，并通过 robot_client 下发 action。",
            "status_fusion 汇总 robot_status、backend_state 和 policy_status，通过 HTTP/WebSocket 更新前端。",
            "停止任务时，/stop 经 task_control 进入 policy_runtime_service 和 robot_client 停止路径。",
        ],
    )

    doc.add_heading("7. 异常与 fallback 原则", level=1)
    add_table(
        doc,
        ["场景", "系统行为", "禁止行为"],
        [
            ["policy 模型未加载", "进入 ERROR 或 paused，记录日志。", "使用 mock_task 替代推理。"],
            ["机器人观测缺失", "暂停闭环，等待 robot_client 恢复。", "生成虚假 action。"],
            ["remote 连接失败", "报告远程 policy runtime 不可用。", "绕过 policy runtime 直接控制硬件。"],
            ["UI 状态源短暂缺失", "status_fusion 可使用 UI fallback 显示。", "让 fallback 进入执行链。"],
        ],
        [Inches(1.7), Inches(2.35), Inches(2.45)],
    )

    doc.add_heading("8. 答辩级验收清单", level=1)
    add_bullets(
        doc,
        [
            "所有文档均描述 frontend -> FastAPI -> task_control -> policy_runtime_service -> robot_client -> SO-ARM101。",
            "local 和 remote 均被定义为同一 AI 控制系统的不同部署方式。",
            "mock_task 在文档中仅作为 UI fallback，不作为控制执行模块。",
            "前端职责限定为任务意图输入和状态展示。",
            "真实硬件动作只通过 robot_client 发送。",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
